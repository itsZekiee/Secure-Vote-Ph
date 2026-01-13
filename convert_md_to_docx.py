import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()

        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)

        # Horizontal Rule
        elif line == '---':
            doc.add_paragraph('' + '_' * 50)

        # Images (Attempt to handle Markdown image syntax ![alt](url))
        elif line.startswith('!['):
            # We can't easily download and embed images without more libs,
            # so we'll just put the image link for now or skip.
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text = match.group(1)
                img_url = match.group(2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"[Image: {alt_text}]")
                doc.add_paragraph(img_url)

        # Lists (unordered)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')

        # Tables (Simple detection)
        elif line.startswith('|'):
            # Basic table support is complex in a simple script,
            # let's just add it as text for now to ensure all content is preserved
            doc.add_paragraph(line)

        # Bold/Italic (Simple regex replacement)
        elif line:
            # Clean up some markdown bold/italic before adding
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
            doc.add_paragraph(clean_line)
        else:
            doc.add_paragraph()

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    markdown_to_docx('DOCUMENTATION.md', 'DOCUMENTATION.docx')
